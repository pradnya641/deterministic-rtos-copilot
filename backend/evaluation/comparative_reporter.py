import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def compute_scores(results):
    scores = {
        "copilot": {
            "compile_passes": 0,
            "continuity_passes": 0,
            "safety_passes": 0,
            "total_violations": 0,
            "rollbacks_triggered": 0,
            "turns_run": 0
        },
        "reference": {
            "compile_passes": 0,
            "continuity_passes": 0,
            "safety_passes": 0,
            "total_violations": 0,
            "rollbacks_triggered": 0,
            "turns_run": 0
        }
    }
    
    for sc in results:
        for t in sc["turns"]:
            # Copilot
            scores["copilot"]["turns_run"] += 1
            if t["copilot"]["compile_ok"]:
                scores["copilot"]["compile_passes"] += 1
            if t["copilot"]["continuity_ok"]:
                scores["copilot"]["continuity_passes"] += 1
            if len(t["copilot"]["violations"]) == 0:
                scores["copilot"]["safety_passes"] += 1
            scores["copilot"]["total_violations"] += len(t["copilot"]["violations"])
            if t["copilot"]["rollback_triggered"]:
                scores["copilot"]["rollbacks_triggered"] += 1
                
            # Reference
            scores["reference"]["turns_run"] += 1
            if t["reference"]["compile_ok"]:
                scores["reference"]["compile_passes"] += 1
            if t["reference"]["continuity_ok"]:
                scores["reference"]["continuity_passes"] += 1
            if len(t["reference"]["violations"]) == 0:
                scores["reference"]["safety_passes"] += 1
            scores["reference"]["total_violations"] += len(t["reference"]["violations"])
            if t["reference"]["rollback_triggered"]:
                scores["reference"]["rollbacks_triggered"] += 1
                
    return scores

def generate_reports(results):
    scores = compute_scores(results)
    
    os.makedirs("evaluation/comparison_results", exist_ok=True)
    
    generate_md_report(results, scores)
    generate_html_report(results, scores)
    generate_docx_report(results, scores)

def generate_md_report(results, scores):
    path = "evaluation/comparison_results/comparison_report.md"
    
    with open(path, "w", encoding="utf-8") as f:
        f.write("# Deterministic Embedded RTOS Copilot — Comparative Evaluation Report\n\n")
        f.write(f"**Date:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        f.write("## 1. Methodology\n\n")
        f.write("Reference responses are captured from real conversational interactions with general-purpose LLMs and replayed deterministically for reproducible comparative evaluation.\n")
        f.write("Both the Deterministic Copilot and the Reference LLM code are extracted using the identical C code parser and compiled under identical constraints using the `arm-none-eabi-gcc` cross-compiler syntax check with FreeRTOS and LPC2148 vendor header stubs.\n\n")
        
        f.write("## 2. Comparative Scores Scorecard\n\n")
        f.write("| Evaluation Dimension | Deterministic Copilot | Reference LLM Corpus |\n")
        f.write("| --- | --- | --- |\n")
        
        c_turns = scores["copilot"]["turns_run"]
        r_turns = scores["reference"]["turns_run"]
        
        f.write(f"| **Compile Success Rate** | {scores['copilot']['compile_passes']}/{c_turns} ({scores['copilot']['compile_passes']/c_turns*100:.1f}%) | {scores['reference']['compile_passes']}/{r_turns} ({scores['reference']['compile_passes']/r_turns*100:.1f}%) |\n")
        f.write(f"| **Mutation Continuity Rate** | {scores['copilot']['continuity_passes']}/{c_turns} ({scores['copilot']['continuity_passes']/c_turns*100:.1f}%) | {scores['reference']['continuity_passes']}/{r_turns} ({scores['reference']['continuity_passes']/r_turns*100:.1f}%) |\n")
        f.write(f"| **RTOS Safety Rate** | {scores['copilot']['safety_passes']}/{c_turns} ({scores['copilot']['safety_passes']/c_turns*100:.1f}%) | {scores['reference']['safety_passes']}/{r_turns} ({scores['reference']['safety_passes']/r_turns*100:.1f}%) |\n")
        f.write(f"| **Rollbacks Triggered** | {scores['copilot']['rollbacks_triggered']} | {scores['reference']['rollbacks_triggered']} |\n")
        f.write(f"| **Total Violations Detected** | {scores['copilot']['total_violations']} | {scores['reference']['total_violations']} |\n\n")
        
        f.write("## 3. Scenario Executions Summary\n\n")
        
        for sc in results:
            f.write(f"### Scenario {sc['scenario']}: {sc['name']}\n\n")
            f.write("| Turn | Query | Copilot status | Reference status | Continuity Delta |\n")
            f.write("| --- | --- | --- | --- | --- |\n")
            
            for t in sc["turns"]:
                c_status = "PASS" if (t["copilot"]["compile_ok"] and len(t["copilot"]["violations"]) == 0) else "FAIL"
                if t["copilot"]["rollback_triggered"]:
                    c_status = "REJECTED (Rollback)"
                r_status = "PASS" if (t["reference"]["compile_ok"] and len(t["reference"]["violations"]) == 0) else "FAIL"
                
                c_lost = ", ".join(t["copilot"]["lost_tasks"]) if t["copilot"]["lost_tasks"] else "None"
                r_lost = ", ".join(t["reference"]["lost_tasks"]) if t["reference"]["lost_tasks"] else "None"
                delta_str = f"Copilot lost: {c_lost} / Ref lost: {r_lost}"
                
                f.write(f"| {t['turn']} | {t['query']} | **{c_status}** | **{r_status}** | {delta_str} |\n")
            f.write("\n")
            
    print(f"Markdown report written to {path}")

def generate_html_report(results, scores):
    path = "evaluation/comparison_results/comparison_report.html"
    
    c_turns = scores["copilot"]["turns_run"]
    r_turns = scores["reference"]["turns_run"]
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Embedded RTOS Copilot — Comparative Benchmark Evaluation</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=Outfit:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-color: #0d0e12;
            --card-bg: #161a23;
            --border-color: #2d3342;
            --text-primary: #f3f4f6;
            --text-secondary: #9ca3af;
            --primary: #4f46e5;
            --primary-light: #818cf8;
            --accent: #06b6d4;
            --success: #10b981;
            --error: #ef4444;
            --warning: #f59e0b;
        }}
        
        body {{
            font-family: 'Inter', sans-serif;
            background-color: var(--bg-color);
            color: var(--text-primary);
            margin: 0;
            padding: 40px;
            line-height: 1.6;
        }}
        
        h1, h2, h3, h4 {{
            font-family: 'Outfit', sans-serif;
            color: #ffffff;
            margin-top: 0;
        }}
        
        h1 {{
            font-size: 2.5rem;
            font-weight: 800;
            background: linear-gradient(90deg, var(--primary-light), var(--accent));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 5px;
        }}
        
        .container {{
            max-width: 1400px;
            margin: 0 auto;
        }}
        
        .header {{
            background: linear-gradient(135deg, #1e2230 0%, #11141e 100%);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 40px;
            margin-bottom: 40px;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.4);
            position: relative;
            overflow: hidden;
        }}
        
        .header::after {{
            content: '';
            position: absolute;
            top: 0;
            right: 0;
            width: 250px;
            height: 250px;
            background: radial-gradient(circle, rgba(79, 70, 229, 0.15) 0%, rgba(0,0,0,0) 70%);
            z-index: 1;
        }}
        
        .header-content {{
            position: relative;
            z-index: 2;
        }}
        
        .methodology {{
            background: rgba(255, 255, 255, 0.03);
            border-left: 4px solid var(--accent);
            padding: 20px;
            border-radius: 4px;
            margin-top: 20px;
            font-size: 0.95rem;
            color: var(--text-secondary);
        }}
        
        .metrics-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(280px, 1fr));
            gap: 25px;
            margin-bottom: 50px;
        }}
        
        .metric-card {{
            background: var(--card-bg);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 25px;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.2);
            transition: transform 0.2s ease, border-color 0.2s ease;
        }}
        
        .metric-card:hover {{
            transform: translateY(-2px);
            border-color: var(--primary-light);
        }}
        
        .metric-title {{
            font-size: 0.85rem;
            color: var(--text-secondary);
            text-transform: uppercase;
            letter-spacing: 1.5px;
            font-weight: 600;
            margin-bottom: 15px;
        }}
        
        .compare-rows {{
            display: flex;
            flex-direction: column;
            gap: 12px;
        }}
        
        .compare-row {{
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        
        .compare-label {{
            font-weight: 500;
            font-size: 0.95rem;
        }}
        
        .compare-val {{
            font-family: 'Outfit', sans-serif;
            font-size: 1.3rem;
            font-weight: 700;
        }}
        
        .copilot-color {{
            color: var(--accent);
        }}
        
        .reference-color {{
            color: var(--warning);
        }}
        
        .scenario-section {{
            background: var(--card-bg);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 35px;
            margin-bottom: 50px;
            box-shadow: 0 4px 25px rgba(0, 0, 0, 0.15);
        }}
        
        .scenario-header {{
            border-bottom: 1px solid var(--border-color);
            padding-bottom: 15px;
            margin-bottom: 25px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }}
        
        .scenario-title {{
            font-size: 1.6rem;
            font-weight: 700;
            margin: 0;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 25px;
        }}
        
        th, td {{
            padding: 14px 18px;
            text-align: left;
            border-bottom: 1px solid var(--border-color);
        }}
        
        th {{
            background-color: rgba(255, 255, 255, 0.02);
            color: var(--text-secondary);
            font-weight: 600;
            font-size: 0.85rem;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        
        .turn-row:hover {{
            background-color: rgba(255, 255, 255, 0.01);
        }}
        
        .badge {{
            padding: 4px 10px;
            border-radius: 20px;
            font-size: 0.75rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            display: inline-block;
        }}
        
        .badge-success {{
            background-color: rgba(16, 185, 129, 0.15);
            color: var(--success);
            border: 1px solid rgba(16, 185, 129, 0.3);
        }}
        
        .badge-error {{
            background-color: rgba(239, 68, 68, 0.15);
            color: var(--error);
            border: 1px solid rgba(239, 68, 68, 0.3);
        }}
        
        .badge-warning {{
            background-color: rgba(245, 158, 11, 0.15);
            color: var(--warning);
            border: 1px solid rgba(245, 158, 11, 0.3);
        }}
        
        .side-by-side {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
            margin-top: 15px;
        }}
        
        .code-box {{
            background-color: #08090d;
            border: 1px solid var(--border-color);
            border-radius: 8px;
            padding: 15px;
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85rem;
            overflow-x: auto;
            max-height: 350px;
            color: #d1d5db;
        }}
        
        .violations-list {{
            margin: 0;
            padding-left: 20px;
            font-size: 0.85rem;
            color: #fca5a5;
        }}
        
        .violations-list.none {{
            color: var(--success);
            list-style: none;
            padding-left: 0;
        }}
        
        .details-btn {{
            background: none;
            border: 1px solid var(--accent);
            color: var(--accent);
            padding: 6px 12px;
            border-radius: 6px;
            cursor: pointer;
            font-size: 0.8rem;
            font-weight: 600;
            transition: all 0.2s ease;
        }}
        
        .details-btn:hover {{
            background-color: var(--accent);
            color: #000;
        }}
        
        .detail-row {{
            display: none;
            background-color: rgba(0, 0, 0, 0.2);
        }}
        
        .expanded {{
            display: table-row !important;
        }}
        
    </style>
    <script>
        function toggleDetail(id) {{
            const row = document.getElementById(id);
            if (row.classList.contains('expanded')) {{
                row.classList.remove('expanded');
            }} else {{
                row.classList.add('expanded');
            }}
        }}
    </script>
</head>
<body>
    <div class="container">
        <div class="header">
            <div class="header-content">
                <h1>RTOS Copilot — Comparative Benchmark Evaluation</h1>
                <p style="color: var(--text-secondary); margin: 0; font-size: 1.1rem;">
                    Hardened verification of deterministic embedded code generation, mutation continuity, and rollback safety.
                </p>
                <div class="methodology">
                    <strong>Evaluation Methodology:</strong> Reference responses are captured from real conversational interactions with general-purpose LLMs and replayed deterministically for reproducible comparative evaluation. Both the Deterministic Copilot and the Reference LLM code are extracted using the identical C code parser and compiled under identical constraints using the <code>arm-none-eabi-gcc</code> cross-compiler syntax check with FreeRTOS and LPC2148 vendor header stubs.
                </div>
            </div>
        </div>

        <div class="metrics-grid">
            <div class="metric-card">
                <div class="metric-title">Compiler Pass Rate</div>
                <div class="compare-rows">
                    <div class="compare-row">
                        <span class="compare-label">Copilot</span>
                        <span class="compare-val copilot-color">{scores['copilot']['compile_passes']}/{c_turns} ({scores['copilot']['compile_passes']/c_turns*100:.1f}%)</span>
                    </div>
                    <div class="compare-row">
                        <span class="compare-label">Reference LLM</span>
                        <span class="compare-val reference-color">{scores['reference']['compile_passes']}/{r_turns} ({scores['reference']['compile_passes']/r_turns*100:.1f}%)</span>
                    </div>
                </div>
            </div>
            
            <div class="metric-card">
                <div class="metric-title">Mutation Continuity Rate</div>
                <div class="compare-rows">
                    <div class="compare-row">
                        <span class="compare-label">Copilot</span>
                        <span class="compare-val copilot-color">{scores['copilot']['continuity_passes']}/{c_turns} ({scores['copilot']['continuity_passes']/c_turns*100:.1f}%)</span>
                    </div>
                    <div class="compare-row">
                        <span class="compare-label">Reference LLM</span>
                        <span class="compare-val reference-color">{scores['reference']['continuity_passes']}/{r_turns} ({scores['reference']['continuity_passes']/r_turns*100:.1f}%)</span>
                    </div>
                </div>
            </div>

            <div class="metric-card">
                <div class="metric-title">RTOS & Platform Safety Rate</div>
                <div class="compare-rows">
                    <div class="compare-row">
                        <span class="compare-label">Copilot</span>
                        <span class="compare-val copilot-color">{scores['copilot']['safety_passes']}/{c_turns} ({scores['copilot']['safety_passes']/c_turns*100:.1f}%)</span>
                    </div>
                    <div class="compare-row">
                        <span class="compare-label">Reference LLM</span>
                        <span class="compare-val reference-color">{scores['reference']['safety_passes']}/{r_turns} ({scores['reference']['safety_passes']/r_turns*100:.1f}%)</span>
                    </div>
                </div>
            </div>
            
            <div class="metric-card">
                <div class="metric-title">Safety Violations & Rollbacks</div>
                <div class="compare-rows">
                    <div class="compare-row">
                        <span class="compare-label">Copilot Violations</span>
                        <span class="compare-val copilot-color" style="color: var(--error);">{scores['copilot']['total_violations']} ({scores['copilot']['rollbacks_triggered']} Rollbacks)</span>
                    </div>
                    <div class="compare-row">
                        <span class="compare-label">Reference Violations</span>
                        <span class="compare-val reference-color" style="color: var(--error);">{scores['reference']['total_violations']} (0 Rollbacks)</span>
                    </div>
                </div>
            </div>
        </div>

        <h2>Detailed Scenario Evaluations</h2>
    """
    
    for sc in results:
        sc_id = sc["scenario"]
        html += f"""
        <div class="scenario-section">
            <div class="scenario-header">
                <h3 class="scenario-title">Scenario {sc_id}: {sc["name"]}</h3>
            </div>
            <table>
                <thead>
                    <tr>
                        <th style="width: 50px;">Turn</th>
                        <th>User Request</th>
                        <th>Copilot Status</th>
                        <th>Reference Status</th>
                        <th>Continuity Comparison</th>
                        <th style="width: 100px;">Actions</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for t in sc["turns"]:
            t_id = t["turn"]
            row_id = f"detail_{sc_id}_{t_id}"
            
            c_comp = t["copilot"]["compile_ok"]
            c_rollback = t["copilot"]["rollback_triggered"]
            c_viols = len(t["copilot"]["violations"])
            
            r_comp = t["reference"]["compile_ok"]
            r_viols = len(t["reference"]["violations"])
            
            # Copilot Badge
            if c_rollback:
                c_badge = f'<span class="badge badge-warning">Rollback ({c_viols} issues)</span>'
            elif c_comp and c_viols == 0:
                c_badge = '<span class="badge badge-success">Pass</span>'
            else:
                c_badge = f'<span class="badge badge-error">Fail ({c_viols} issues)</span>'
                
            # Reference Badge
            if r_comp and r_viols == 0:
                r_badge = '<span class="badge badge-success">Pass</span>'
            else:
                r_badge = f'<span class="badge badge-error">Fail ({r_viols} issues)</span>'
                
            c_tasks_str = ", ".join(t["copilot"]["tasks"]) if t["copilot"]["tasks"] else "None"
            r_tasks_str = ", ".join(t["reference"]["tasks"]) if t["reference"]["tasks"] else "None"
            
            continuity_str = f"""
                <div style="font-size: 0.85rem;">
                    <strong>Copilot Tasks:</strong> <code style="color: var(--accent);">{c_tasks_str}</code>
                    {" " if t["copilot"]["continuity_ok"] else '<span style="color: var(--error); font-weight: bold;">(State Lost!)</span>'}
                    <br>
                    <strong>Reference Tasks:</strong> <code style="color: var(--warning);">{r_tasks_str}</code>
                    {" " if t["reference"]["continuity_ok"] else '<span style="color: var(--error); font-weight: bold;">(State Lost!)</span>'}
                </div>
            """
            
            html += f"""
                    <tr class="turn-row">
                        <td>{t_id}</td>
                        <td style="font-weight: 500; max-width: 300px;">{t["query"]}</td>
                        <td>{c_badge}</td>
                        <td>{r_badge}</td>
                        <td>{continuity_str}</td>
                        <td>
                            <button class="details-btn" onclick="toggleDetail('{row_id}')">Compare Code</button>
                        </td>
                    </tr>
                    <tr id="{row_id}" class="detail-row">
                        <td colspan="6" style="padding: 20px; background-color: rgba(0,0,0,0.15);">
                            <div class="side-by-side">
                                <div>
                                    <h4 style="color: var(--accent); border-bottom: 1px solid var(--border-color); padding-bottom: 5px;">Copilot (Staged & Validated)</h4>
                                    <strong>Compiler Status:</strong> {t["copilot"]["compile_msg"]}<br>
                                    <strong>Detected RTOS/Safety Violations:</strong>
                                    {f'<ul class="violations-list">{"".join(f"<li>{v}</li>" for v in t["copilot"]["violations"])}</ul>' if t["copilot"]["violations"] else '<ul class="violations-list none"><li>✔ None</li></ul>'}
                                    <div style="margin-top: 10px;"><strong>Extracted Code:</strong></div>
                                    <pre class="code-box"><code>{t["copilot"]["code"] if t["copilot"]["code"] else "/* No Code Blocks */"}</code></pre>
                                </div>
                                <div>
                                    <h4 style="color: var(--warning); border-bottom: 1px solid var(--border-color); padding-bottom: 5px;">Reference LLM Response</h4>
                                    <strong>Compiler Status:</strong> {t["reference"]["compile_msg"]}<br>
                                    <strong>Detected RTOS/Safety Violations:</strong>
                                    {f'<ul class="violations-list">{"".join(f"<li>{v}</li>" for v in t["reference"]["violations"])}</ul>' if t["reference"]["violations"] else '<ul class="violations-list none"><li>✔ None</li></ul>'}
                                    <div style="margin-top: 10px;"><strong>Extracted Code:</strong></div>
                                    <pre class="code-box"><code>{t["reference"]["code"] if t["reference"]["code"] else "/* No Code Blocks */"}</code></pre>
                                </div>
                            </div>
                        </td>
                    </tr>
            """
        html += """
                </tbody>
            </table>
        </div>
        """
        
    html += """
    </div>
</body>
</html>
    """
    
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"HTML report written to {path}")

def generate_docx_report(results, scores):
    path = "evaluation/comparison_results/comparison_report.docx"
    doc = Document()
    
    # Page setup - 0.75 margin for more space
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styling helpers
    c_blue = RGBColor(16, 44, 87)
    c_gray = RGBColor(100, 100, 100)
    c_success = RGBColor(16, 185, 129)
    c_error = RGBColor(239, 68, 68)
    
    # Title
    t_para = doc.add_paragraph()
    t_run = t_para.add_run("RTOS Copilot — Comparative Benchmark Evaluation")
    t_run.font.name = 'Outfit'
    t_run.font.size = Pt(22)
    t_run.font.bold = True
    t_run.font.color.rgb = c_blue
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    d_para = doc.add_paragraph()
    d_para.add_run(f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    d_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Methodology
    doc.add_heading("1. Methodology & Fair Sandbox Setup", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Reference responses are captured from real conversational interactions with general-purpose LLMs and replayed "
        "deterministically for reproducible comparative evaluation. Both the Deterministic Copilot and the Reference LLM "
        "code are extracted using the identical C code parser and compiled under identical constraints using the "
        "arm-none-eabi-gcc cross-compiler syntax check with FreeRTOS and LPC2148 vendor header stubs. This ensures 100% fair "
        "conditions for compilation rate checks."
    )
    
    # Metrics
    doc.add_heading("2. Executive Scorecard Summary", level=1)
    
    c_turns = scores["copilot"]["turns_run"]
    r_turns = scores["reference"]["turns_run"]
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Set headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Evaluation Dimension'
    hdr_cells[1].text = 'Deterministic Copilot'
    hdr_cells[2].text = 'Reference LLM Corpus'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    metrics_data = [
        ("Compile Success Rate", f"{scores['copilot']['compile_passes']}/{c_turns} ({scores['copilot']['compile_passes']/c_turns*100:.1f}%)", f"{scores['reference']['compile_passes']}/{r_turns} ({scores['reference']['compile_passes']/r_turns*100:.1f}%)"),
        ("Mutation Continuity Rate", f"{scores['copilot']['continuity_passes']}/{c_turns} ({scores['copilot']['continuity_passes']/c_turns*100:.1f}%)", f"{scores['reference']['continuity_passes']}/{r_turns} ({scores['reference']['continuity_passes']/r_turns*100:.1f}%)"),
        ("RTOS & Platform Safety Rate", f"{scores['copilot']['safety_passes']}/{c_turns} ({scores['copilot']['safety_passes']/c_turns*100:.1f}%)", f"{scores['reference']['safety_passes']}/{r_turns} ({scores['reference']['safety_passes']/r_turns*100:.1f}%)"),
        ("Rollbacks Triggered", f"{scores['copilot']['rollbacks_triggered']}", f"{scores['reference']['rollbacks_triggered']}"),
        ("Total Safety Violations", f"{scores['copilot']['total_violations']}", f"{scores['reference']['total_violations']}")
    ]
    
    for idx, (dim, c_val, r_val) in enumerate(metrics_data, 1):
        row = table.rows[idx]
        row.cells[0].text = dim
        row.cells[1].text = c_val
        row.cells[2].text = r_val
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        
    doc.add_paragraph()
    
    # Detailed Scenarios
    doc.add_heading("3. Detailed Scenario Evaluations", level=1)
    
    for sc in results:
        doc.add_heading(f"Scenario {sc['scenario']}: {sc['name']}", level=2)
        
        sc_table = doc.add_table(rows=1, cols=5)
        sc_table.style = 'Light Shading Accent 1'
        hdr_cells = sc_table.rows[0].cells
        hdr_cells[0].text = 'Turn'
        hdr_cells[1].text = 'User Query'
        hdr_cells[2].text = 'Copilot Status'
        hdr_cells[3].text = 'Reference Status'
        hdr_cells[4].text = 'Continuity Delta'
        
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            
        for t in sc["turns"]:
            row_cells = sc_table.add_row().cells
            row_cells[0].text = str(t["turn"])
            row_cells[1].text = t["query"]
            
            c_comp = t["copilot"]["compile_ok"]
            c_rollback = t["copilot"]["rollback_triggered"]
            c_viols = len(t["copilot"]["violations"])
            
            r_comp = t["reference"]["compile_ok"]
            r_viols = len(t["reference"]["violations"])
            
            if c_rollback:
                row_cells[2].text = f"REJECTED\n({c_viols} issues)"
            elif c_comp and c_viols == 0:
                row_cells[2].text = "PASS"
            else:
                row_cells[2].text = f"FAIL\n({c_viols} issues)"
                
            if r_comp and r_viols == 0:
                row_cells[3].text = "PASS"
            else:
                row_cells[3].text = f"FAIL\n({r_viols} issues)"
                
            c_lost = ", ".join(t["copilot"]["lost_tasks"]) if t["copilot"]["lost_tasks"] else "None"
            r_lost = ", ".join(t["reference"]["lost_tasks"]) if t["reference"]["lost_tasks"] else "None"
            row_cells[4].text = f"Copilot lost: {c_lost}\nRef lost: {r_lost}"
            
        doc.add_paragraph()
        
    doc.save(path)
    print(f"DOCX report written to {path}")
