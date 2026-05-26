import os
import json
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_reports(all_results, compare_file=None):
    print("Generating evaluation reports...")
    
    # 1. Load comparison results if any
    compare_data = None
    if compare_file and os.path.exists(compare_file):
        try:
            with open(compare_file, "r", encoding="utf-8") as f:
                compare_data = json.load(f)
            print(f"Loaded comparison log from {compare_file}")
        except Exception as e:
            print(f"Could not load comparison file: {e}")

    # 2. Compute metrics
    metrics = compute_metrics(all_results)
    compare_metrics = compute_metrics(compare_data) if compare_data else None

    # 3. Generate Markdown Report
    generate_md_report(metrics, all_results, compare_metrics)
    
    # 4. Generate HTML Report
    generate_html_report(metrics, all_results, compare_metrics)
    
    # 5. Generate DOCX Report
    generate_docx_report(metrics, all_results, compare_metrics)

def compute_metrics(results):
    if not results:
        return None
    
    total_turns = 0
    compile_passes = 0
    total_compiles = 0
    rollbacks = 0
    total_latency = 0
    failures = 0
    
    for suite, scenarios in results.items():
        for scenario in scenarios:
            for turn in scenario.get("turns", []):
                total_turns += 1
                total_latency += turn.get("latency_ms", 0)
                if turn.get("rollback_triggered"):
                    rollbacks += 1
                if turn.get("validation_errors") or not turn.get("compile_valid"):
                    failures += 1
                
                # Check compile status
                c_code = turn.get("response", "")
                if "#include" in c_code:
                    total_compiles += 1
                    if turn.get("compile_valid"):
                        compile_passes += 1
                        
    avg_latency = total_latency / total_turns if total_turns > 0 else 0
    compile_rate = (compile_passes / total_compiles * 100) if total_compiles > 0 else 100.0
    
    return {
        "total_turns": total_turns,
        "compile_passes": compile_passes,
        "total_compiles": total_compiles,
        "compile_rate": compile_rate,
        "rollbacks": rollbacks,
        "avg_latency_ms": avg_latency,
        "failures": failures,
        "success_rate": ((total_turns - failures) / total_turns * 100) if total_turns > 0 else 100.0
    }

def generate_md_report(metrics, results, compare_metrics):
    path = "evaluation/results/report.md"
    with open(path, "w", encoding="utf-8") as f:
        f.write("# Deterministic Embedded RTOS Copilot — Stability & Evaluation Report\n\n")
        f.write(f"**Date:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        # Summary table
        f.write("## Executive Metrics Summary\n\n")
        f.write("| Metric | Current Value | Baseline | Delta |\n")
        f.write("| --- | --- | --- | --- |\n")
        
        def write_metric_row(name, key, fmt_func):
            curr = fmt_func(metrics.get(key))
            base = fmt_func(compare_metrics.get(key)) if compare_metrics else "N/A"
            if compare_metrics:
                diff_val = metrics.get(key) - compare_metrics.get(key)
                delta = f"{'+' if diff_val > 0 else ''}{fmt_func(diff_val)}"
            else:
                delta = "N/A"
            f.write(f"| {name} | {curr} | {base} | {delta} |\n")
            
        write_metric_row("Total Turns Run", "total_turns", lambda x: f"{x}" if x is not None else "")
        write_metric_row("Success Rate (%)", "success_rate", lambda x: f"{x:.1f}%" if x is not None else "")
        write_metric_row("Compiler Pass Rate (%)", "compile_rate", lambda x: f"{x:.1f}%" if x is not None else "")
        write_metric_row("Rollbacks Triggered", "rollbacks", lambda x: f"{x}" if x is not None else "")
        write_metric_row("Average Latency (ms)", "avg_latency_ms", lambda x: f"{int(x)}ms" if x is not None else "")
        
        f.write("\n\n## Detailed Test Suite Results\n\n")
        for suite, scenarios in results.items():
            f.write(f"### Suite: {suite.upper()}\n\n")
            for sc in scenarios:
                f.write(f"#### Scenario: {sc['scenario_name']}\n\n")
                f.write("| Turn | Query | Status | Compiler | Rollback | Latency |\n")
                f.write("| --- | --- | --- | --- | --- | --- |\n")
                for t in sc["turns"]:
                    status = "FAIL" if (t["validation_errors"] or not t["compile_valid"]) else "PASS"
                    comp = "OK" if t["compile_valid"] else "FAILED"
                    roll = "YES" if t["rollback_triggered"] else "NO"
                    f.write(f"| {t['turn']} | {t['query']} | **{status}** | {comp} | {roll} | {t['latency_ms']}ms |\n")
                f.write("\n")
                
    print(f"Markdown report written to {path}")

def generate_html_report(metrics, results, compare_metrics):
    path = "evaluation/results/report.html"
    
    # Harmony HSL Colors / sleek dark theme style
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Embedded RTOS Copilot Stability Report</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&family=Outfit:wght@300;400;600;700&display=swap" rel="stylesheet">
    <style>
        body {{
            font-family: 'Inter', sans-serif;
            background-color: #0d0e12;
            color: #e2e4e9;
            margin: 0;
            padding: 40px;
        }}
        h1, h2, h3, h4 {{
            font-family: 'Outfit', sans-serif;
            color: #ffffff;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
        }}
        .header {{
            background: linear-gradient(135deg, #1f242e 0%, #11141a 100%);
            border: 1px solid #2d3342;
            border-radius: 12px;
            padding: 30px;
            margin-bottom: 30px;
            box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.3);
        }}
        .header h1 {{
            margin: 0 0 10px 0;
            font-size: 2.2rem;
            background: linear-gradient(90deg, #4f46e5, #06b6d4);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}
        .metrics-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
            gap: 20px;
            margin-bottom: 40px;
        }}
        .metric-card {{
            background: #161a23;
            border: 1px solid #2d3342;
            border-radius: 8px;
            padding: 20px;
            text-align: center;
        }}
        .metric-val {{
            font-size: 2rem;
            font-weight: 700;
            color: #06b6d4;
            margin-top: 5px;
        }}
        .metric-label {{
            font-size: 0.9rem;
            color: #94a3b8;
            text-transform: uppercase;
            letter-spacing: 1px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 30px;
            background: #11141a;
            border-radius: 8px;
            overflow: hidden;
            border: 1px solid #2d3342;
        }}
        th, td {{
            padding: 14px 20px;
            text-align: left;
            border-bottom: 1px solid #2d3342;
        }}
        th {{
            background-color: #1a1f2c;
            color: #94a3b8;
            font-weight: 600;
        }}
        .status-pass {{
            color: #10b981;
            font-weight: bold;
        }}
        .status-fail {{
            color: #ef4444;
            font-weight: bold;
        }}
        .suite-section {{
            background: #161a23;
            border: 1px solid #2d3342;
            border-radius: 8px;
            padding: 25px;
            margin-bottom: 35px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>RTOS Copilot Stability Report</h1>
            <p style="color: #94a3b8; margin: 0;">Date generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>

        <div class="metrics-grid">
            <div class="metric-card">
                <div class="metric-label">Success Rate</div>
                <div class="metric-val">{metrics['success_rate']:.1f}%</div>
            </div>
            <div class="metric-card">
                <div class="metric-label">Compiler Pass Rate</div>
                <div class="metric-val">{metrics['compile_rate']:.1f}%</div>
            </div>
            <div class="metric-card">
                <div class="metric-label">Rollbacks</div>
                <div class="metric-val">{metrics['rollbacks']}</div>
            </div>
            <div class="metric-card">
                <div class="metric-label">Avg Latency</div>
                <div class="metric-val">{int(metrics['avg_latency_ms'])}ms</div>
            </div>
        </div>

        <h2>Detailed Test Suites</h2>
    """

    for suite, scenarios in results.items():
        html_content += f"""
        <div class="suite-section">
            <h3>Suite: {suite.upper()}</h3>
        """
        for sc in scenarios:
            html_content += f"""
            <h4>Scenario: {sc['scenario_name']}</h4>
            <table>
                <thead>
                    <tr>
                        <th>Turn</th>
                        <th>Query</th>
                        <th>Status</th>
                        <th>Compiler</th>
                        <th>Rollback</th>
                        <th>Latency</th>
                    </tr>
                </thead>
                <tbody>
            """
            for t in sc["turns"]:
                status_class = "status-pass" if not (t["validation_errors"] or not t["compile_valid"]) else "status-fail"
                status_text = "PASS" if not (t["validation_errors"] or not t["compile_valid"]) else "FAIL"
                comp_text = "OK" if t["compile_valid"] else "FAILED"
                roll_text = "YES" if t["rollback_triggered"] else "NO"
                html_content += f"""
                    <tr>
                        <td>{t['turn']}</td>
                        <td>{t['query']}</td>
                        <td class="{status_class}">{status_text}</td>
                        <td>{comp_text}</td>
                        <td>{roll_text}</td>
                        <td>{t['latency_ms']}ms</td>
                    </tr>
                """
            html_content += """
                </tbody>
            </table>
            """
        html_content += "</div>"

    html_content += """
    </div>
</body>
</html>
    """
    
    with open(path, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"HTML report written to {path}")

def generate_docx_report(metrics, results, compare_metrics):
    path = "evaluation/results/report.docx"
    doc = Document()
    
    # Theme colors
    primary_color = RGBColor(31, 78, 121)   # Blue
    secondary_color = RGBColor(80, 80, 80)  # Gray
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run("RTOS Copilot Stability & Evaluation Report")
    run.font.name = 'Outfit'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = primary_color
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_p = doc.add_paragraph()
    date_p.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Summary header
    h = doc.add_heading("1. Executive Summary Metrics", level=1)
    for run in h.runs:
        run.font.color.rgb = primary_color
        
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Shading Accent 1'
    
    summary_data = [
        ("Total Turns Run", f"{metrics['total_turns']}"),
        ("Success Rate (%)", f"{metrics['success_rate']:.1f}%"),
        ("Compiler Pass Rate (%)", f"{metrics['compile_rate']:.1f}%"),
        ("Rollbacks Triggered", f"{metrics['rollbacks']}"),
        ("Average Latency (ms)", f"{int(metrics['avg_latency_ms'])}ms")
    ]
    
    for idx, (label, val) in enumerate(summary_data):
        row = table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = val
        
    doc.add_paragraph()
    
    # Detailed section
    h2 = doc.add_heading("2. Detailed Test Suite Execution", level=1)
    for run in h2.runs:
        run.font.color.rgb = primary_color
        
    for suite, scenarios in results.items():
        doc.add_heading(f"Suite: {suite.upper()}", level=2)
        for sc in scenarios:
            doc.add_heading(f"Scenario: {sc['scenario_name']}", level=3)
            
            # Scenario table
            sc_table = doc.add_table(rows=1, cols=6)
            sc_table.style = 'Light Shading Accent 1'
            hdr_cells = sc_table.rows[0].cells
            hdr_cells[0].text = 'Turn'
            hdr_cells[1].text = 'Query'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Compiler'
            hdr_cells[4].text = 'Rollback'
            hdr_cells[5].text = 'Latency'
            
            for t in sc["turns"]:
                row_cells = sc_table.add_row().cells
                row_cells[0].text = str(t["turn"])
                row_cells[1].text = t["query"]
                row_cells[2].text = "PASS" if not (t["validation_errors"] or not t["compile_valid"]) else "FAIL"
                row_cells[3].text = "OK" if t["compile_valid"] else "FAILED"
                row_cells[4].text = "YES" if t["rollback_triggered"] else "NO"
                row_cells[5].text = f"{t['latency_ms']}ms"
                
            doc.add_paragraph()

    doc.save(path)
    print(f"DOCX report written to {path}")
